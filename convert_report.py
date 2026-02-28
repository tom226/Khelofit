"""
Convert KHELOFIT_PROJECT_REPORT.md to DOCX and PDF formats.
"""
import re
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch, cm
from reportlab.lib.colors import HexColor, black, white, grey
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    PageBreak, KeepTogether, HRFlowable
)

REPORT_PATH = os.path.join(os.path.dirname(__file__), "KHELOFIT_PROJECT_REPORT.md")
DOCX_PATH = os.path.join(os.path.dirname(__file__), "KHELOFIT_PROJECT_REPORT.docx")
PDF_PATH = os.path.join(os.path.dirname(__file__), "KHELOFIT_PROJECT_REPORT.pdf")

# ── Read markdown ──
with open(REPORT_PATH, "r", encoding="utf-8") as f:
    md_lines = f.readlines()

# ── Helpers ──
def strip_md(text):
    """Remove markdown formatting from text."""
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'`(.+?)`', r'\1', text)
    text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
    text = text.replace('\\', '')
    return text.strip()

def parse_table(lines, start_idx):
    """Parse markdown table starting at start_idx. Returns (rows, end_idx)."""
    rows = []
    i = start_idx
    while i < len(lines):
        line = lines[i].strip()
        if not line.startswith('|'):
            break
        cells = [c.strip() for c in line.split('|')[1:-1]]
        # Skip separator row
        if cells and all(re.match(r'^[-:]+$', c) for c in cells):
            i += 1
            continue
        rows.append(cells)
        i += 1
    return rows, i


# ====================================================================
# PART 1: Generate DOCX
# ====================================================================
print("Generating DOCX...")
doc = Document()

# Page setup
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# Title style
style = doc.styles['Title']
style.font.size = Pt(28)
style.font.color.rgb = RGBColor(0x00, 0x4E, 0x89)
style.font.bold = True

style = doc.styles['Heading 1']
style.font.size = Pt(18)
style.font.color.rgb = RGBColor(0xFF, 0x6B, 0x35)
style.font.bold = True

style = doc.styles['Heading 2']
style.font.size = Pt(14)
style.font.color.rgb = RGBColor(0x00, 0x4E, 0x89)
style.font.bold = True

style = doc.styles['Heading 3']
style.font.size = Pt(12)
style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.font.bold = True

# Normal text
style = doc.styles['Normal']
style.font.size = Pt(10)
style.font.name = 'Calibri'
style.paragraph_format.space_after = Pt(4)

def add_formatted_paragraph(doc, text, style_name='Normal'):
    """Add paragraph with bold/italic inline formatting."""
    p = doc.add_paragraph(style=style_name)
    # Split by bold markers
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(strip_md(part))
            run.bold = True
        else:
            # Handle italic
            sub_parts = re.split(r'(\*.*?\*)', part)
            for sp in sub_parts:
                if sp.startswith('*') and sp.endswith('*') and not sp.startswith('**'):
                    run = p.add_run(strip_md(sp))
                    run.italic = True
                else:
                    p.add_run(strip_md(sp))
    return p

def add_docx_table(doc, rows):
    """Add a formatted table to the document."""
    if not rows:
        return
    num_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            if j < num_cols:
                cell = table.cell(i, j)
                cell.text = strip_md(cell_text)
                for paragraph in cell.paragraphs:
                    paragraph.style.font.size = Pt(9)
                    if i == 0:
                        for run in paragraph.runs:
                            run.bold = True

    doc.add_paragraph()

# Process markdown lines
i = 0
in_code_block = False

while i < len(md_lines):
    line = md_lines[i]
    stripped = line.strip()

    # Skip empty lines
    if not stripped:
        i += 1
        continue

    # Code blocks
    if stripped.startswith('```'):
        in_code_block = not in_code_block
        i += 1
        continue
    if in_code_block:
        p = doc.add_paragraph(stripped, style='Normal')
        p.style.font.name = 'Consolas'
        i += 1
        continue

    # Horizontal rules
    if stripped == '---':
        doc.add_paragraph()
        i += 1
        continue

    # Title (# )
    if stripped.startswith('# ') and not stripped.startswith('## '):
        text = strip_md(stripped[2:])
        doc.add_paragraph(text, style='Title')
        i += 1
        continue

    # H2 (## )
    if stripped.startswith('## ') and not stripped.startswith('### '):
        text = strip_md(stripped[3:])
        doc.add_heading(text, level=1)
        i += 1
        continue

    # H3 (### )
    if stripped.startswith('### ') and not stripped.startswith('#### '):
        text = strip_md(stripped[4:])
        doc.add_heading(text, level=2)
        i += 1
        continue

    # H4 (#### )
    if stripped.startswith('#### '):
        text = strip_md(stripped[5:])
        doc.add_heading(text, level=3)
        i += 1
        continue

    # Tables
    if stripped.startswith('|'):
        rows, end_i = parse_table(md_lines, i)
        add_docx_table(doc, rows)
        i = end_i
        continue

    # Bullet lists
    if stripped.startswith('- '):
        text = strip_md(stripped[2:])
        add_formatted_paragraph(doc, text, 'List Bullet')
        i += 1
        continue

    # Checkbox lists
    if stripped.startswith('- [ ]') or stripped.startswith('- [x]'):
        checked = stripped.startswith('- [x]')
        text = strip_md(stripped[6:])
        prefix = "☑ " if checked else "☐ "
        doc.add_paragraph(prefix + text, style='List Bullet')
        i += 1
        continue

    # Blockquotes
    if stripped.startswith('> '):
        text = strip_md(stripped[2:])
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        run = p.add_run(text)
        run.italic = True
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        i += 1
        continue

    # Regular paragraph
    if stripped:
        add_formatted_paragraph(doc, stripped)

    i += 1

doc.save(DOCX_PATH)
print(f"✅ DOCX saved: {DOCX_PATH}")


# ====================================================================
# PART 2: Generate PDF
# ====================================================================
print("Generating PDF...")

pdf_doc = SimpleDocTemplate(
    PDF_PATH,
    pagesize=A4,
    topMargin=1.5*cm,
    bottomMargin=1.5*cm,
    leftMargin=2*cm,
    rightMargin=2*cm,
    title="KheloFit Project Report",
    author="KheloFit"
)

styles = getSampleStyleSheet()

# Custom styles
styles.add(ParagraphStyle(
    'DocTitle',
    parent=styles['Title'],
    fontSize=26,
    textColor=HexColor('#004E89'),
    spaceAfter=6,
    alignment=TA_CENTER,
    fontName='Helvetica-Bold'
))

styles.add(ParagraphStyle(
    'DocSubtitle',
    parent=styles['Normal'],
    fontSize=12,
    textColor=HexColor('#FF6B35'),
    spaceAfter=4,
    alignment=TA_CENTER,
    fontName='Helvetica-Bold'
))

styles.add(ParagraphStyle(
    'SectionHead',
    parent=styles['Heading1'],
    fontSize=16,
    textColor=HexColor('#FF6B35'),
    spaceBefore=18,
    spaceAfter=8,
    fontName='Helvetica-Bold',
    leftIndent=0
))

styles.add(ParagraphStyle(
    'SubSectionHead',
    parent=styles['Heading2'],
    fontSize=13,
    textColor=HexColor('#004E89'),
    spaceBefore=12,
    spaceAfter=6,
    fontName='Helvetica-Bold',
    leftIndent=0
))

styles.add(ParagraphStyle(
    'SubSubSectionHead',
    parent=styles['Heading3'],
    fontSize=11,
    textColor=HexColor('#333333'),
    spaceBefore=8,
    spaceAfter=4,
    fontName='Helvetica-Bold',
    leftIndent=0
))

styles.add(ParagraphStyle(
    'BodyText2',
    parent=styles['Normal'],
    fontSize=9.5,
    leading=13,
    spaceAfter=4,
    fontName='Helvetica',
    alignment=TA_JUSTIFY
))

styles.add(ParagraphStyle(
    'BulletItem',
    parent=styles['Normal'],
    fontSize=9.5,
    leading=13,
    spaceAfter=2,
    leftIndent=20,
    bulletIndent=10,
    fontName='Helvetica'
))

styles.add(ParagraphStyle(
    'BlockQuote',
    parent=styles['Normal'],
    fontSize=9.5,
    leading=13,
    spaceAfter=6,
    leftIndent=30,
    fontName='Helvetica-Oblique',
    textColor=HexColor('#555555')
))

def md_to_pdf_text(text):
    """Convert markdown inline formatting to reportlab XML tags."""
    text = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', text)
    text = re.sub(r'\*(.+?)\*', r'<i>\1</i>', text)
    text = re.sub(r'`(.+?)`', r'<font face="Courier" size="9">\1</font>', text)
    text = re.sub(r'\[(.+?)\]\(.+?\)', r'<u>\1</u>', text)
    # Escape special XML chars that aren't already tags
    text = text.replace('&', '&amp;')
    # Re-fix our tags that got broken
    text = text.replace('&amp;amp;', '&amp;')
    return text

def make_pdf_table(rows):
    """Create a reportlab Table from parsed rows."""
    if not rows:
        return None

    # Calculate col widths based on page
    available = A4[0] - 4*cm
    num_cols = len(rows[0])
    col_width = available / num_cols

    data = []
    for row in rows:
        cells = []
        for cell in row:
            cell_text = strip_md(cell)
            cells.append(Paragraph(cell_text, styles['BodyText2']))
        # Pad if needed
        while len(cells) < num_cols:
            cells.append(Paragraph('', styles['BodyText2']))
        data.append(cells)

    if not data:
        return None

    t = Table(data, colWidths=[col_width]*num_cols, repeatRows=1)
    
    style_cmds = [
        ('BACKGROUND', (0, 0), (-1, 0), HexColor('#004E89')),
        ('TEXTCOLOR', (0, 0), (-1, 0), white),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 8),
        ('FONTSIZE', (0, 0), (-1, 0), 8.5),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('GRID', (0, 0), (-1, -1), 0.5, HexColor('#CCCCCC')),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [white, HexColor('#F5F5F5')]),
        ('TOPPADDING', (0, 0), (-1, -1), 4),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
        ('LEFTPADDING', (0, 0), (-1, -1), 6),
        ('RIGHTPADDING', (0, 0), (-1, -1), 6),
    ]
    t.setStyle(TableStyle(style_cmds))
    return t


# Build PDF content
elements = []

# Cover page
elements.append(Spacer(1, 3*cm))
elements.append(Paragraph("KheloFit", styles['DocTitle']))
elements.append(Spacer(1, 0.5*cm))
elements.append(Paragraph("India's AI Health + Sports + Events Super App", styles['DocSubtitle']))
elements.append(Spacer(1, 0.3*cm))
elements.append(HRFlowable(width="60%", thickness=2, color=HexColor('#FF6B35'), spaceAfter=10))
elements.append(Spacer(1, 0.5*cm))
elements.append(Paragraph("Complete Project Report", styles['DocSubtitle']))
elements.append(Paragraph("Marketing & Development Plan", styles['DocSubtitle']))
elements.append(Paragraph("Target: $1,000,000 Revenue", styles['DocSubtitle']))
elements.append(Spacer(1, 1*cm))
elements.append(Paragraph("February 2026 | Version 1.0", ParagraphStyle(
    'CoverDate', parent=styles['Normal'], fontSize=10, alignment=TA_CENTER,
    textColor=HexColor('#888888')
)))
elements.append(Paragraph("Confidential - KheloFit Internal Use Only", ParagraphStyle(
    'CoverConf', parent=styles['Normal'], fontSize=9, alignment=TA_CENTER,
    textColor=HexColor('#CC0000')
)))
elements.append(PageBreak())

# Process content
i = 0
in_code_block = False
skip_title = True  # Skip the first title (we made a cover page)

while i < len(md_lines):
    line = md_lines[i]
    stripped = line.strip()

    if not stripped:
        i += 1
        continue

    # Code blocks
    if stripped.startswith('```'):
        in_code_block = not in_code_block
        i += 1
        continue
    if in_code_block:
        elements.append(Paragraph(
            stripped.replace('<', '&lt;').replace('>', '&gt;'),
            ParagraphStyle('Code', parent=styles['Normal'], fontSize=8,
                          fontName='Courier', leftIndent=20,
                          backColor=HexColor('#F5F5F5'), spaceAfter=1)
        ))
        i += 1
        continue

    # Horizontal rules
    if stripped == '---':
        elements.append(HRFlowable(width="100%", thickness=1, color=HexColor('#DDDDDD'), spaceAfter=8, spaceBefore=8))
        i += 1
        continue

    # Title
    if stripped.startswith('# ') and not stripped.startswith('## '):
        if skip_title:
            skip_title = False
            i += 1
            continue
        text = strip_md(stripped[2:])
        elements.append(Paragraph(text, styles['DocTitle']))
        i += 1
        continue

    # H2
    if stripped.startswith('## ') and not stripped.startswith('### '):
        text = strip_md(stripped[3:])
        elements.append(Paragraph(text, styles['SectionHead']))
        i += 1
        continue

    # H3
    if stripped.startswith('### ') and not stripped.startswith('#### '):
        text = strip_md(stripped[4:])
        elements.append(Paragraph(text, styles['SubSectionHead']))
        i += 1
        continue

    # H4
    if stripped.startswith('#### '):
        text = strip_md(stripped[5:])
        elements.append(Paragraph(text, styles['SubSubSectionHead']))
        i += 1
        continue

    # Tables
    if stripped.startswith('|'):
        rows, end_i = parse_table(md_lines, i)
        tbl = make_pdf_table(rows)
        if tbl:
            elements.append(tbl)
            elements.append(Spacer(1, 0.3*cm))
        i = end_i
        continue

    # Bullet lists
    if stripped.startswith('- [ ]') or stripped.startswith('- [x]'):
        checked = stripped.startswith('- [x]')
        text = strip_md(stripped[6:])
        prefix = "☑ " if checked else "☐ "
        elements.append(Paragraph(prefix + text, styles['BulletItem']))
        i += 1
        continue

    if stripped.startswith('- '):
        text = md_to_pdf_text(stripped[2:])
        elements.append(Paragraph("• " + text, styles['BulletItem']))
        i += 1
        continue

    # Blockquotes
    if stripped.startswith('> '):
        text = md_to_pdf_text(stripped[2:])
        elements.append(Paragraph(text, styles['BlockQuote']))
        i += 1
        continue

    # Regular paragraph
    if stripped:
        text = md_to_pdf_text(stripped)
        try:
            elements.append(Paragraph(text, styles['BodyText2']))
        except Exception:
            elements.append(Paragraph(strip_md(stripped), styles['BodyText2']))

    i += 1

# Page numbers
def add_page_number(canvas, doc):
    canvas.saveState()
    canvas.setFont('Helvetica', 8)
    canvas.setFillColor(HexColor('#888888'))
    page_num = canvas.getPageNumber()
    text = f"KheloFit Project Report  |  Page {page_num}"
    canvas.drawCentredString(A4[0]/2, 1*cm, text)
    canvas.restoreState()

pdf_doc.build(elements, onFirstPage=add_page_number, onLaterPages=add_page_number)
print(f"✅ PDF saved: {PDF_PATH}")
print("\nDone! Both files generated successfully.")
